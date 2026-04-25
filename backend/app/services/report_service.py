from __future__ import annotations

import io
from datetime import date, datetime, timezone

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree
from sqlalchemy import func, or_, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.ops_models import InventoryItem, Product
from app.services.ventas_service import VentasService

# ── Paleta de colores ──────────────────────────────────────────────────────────
_BRAND_BLUE = RGBColor(0x1E, 0x40, 0xAF)
_BRAND_BLUE_LIGHT = RGBColor(0xDB, 0xEA, 0xFE)
_DARK_TEXT = RGBColor(0x1F, 0x29, 0x37)
_GRAY = RGBColor(0x6B, 0x72, 0x80)
_LIGHT_GRAY = RGBColor(0xF3, 0xF4, 0xF6)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Colores semánticos
_RED = RGBColor(0xDC, 0x26, 0x26)
_RED_BG = "FEE2E2"
_ORANGE = RGBColor(0xEA, 0x58, 0x0C)
_ORANGE_BG = "FFEDD5"
_YELLOW = RGBColor(0xCA, 0x8A, 0x04)
_YELLOW_BG = "FEF9C3"
_GREEN = RGBColor(0x16, 0xA3, 0x4A)
_GREEN_BG = "DCFCE7"

_SECTION_KEYS = {"kpis", "clientes", "productos", "margen", "pagos", "riesgo"}
_ALMACEN_SECTION_KEYS = {"kpis", "valor", "alertas", "dormidos"}


def _fmt_mxn(value: float | None) -> str:
    if value is None:
        return "—"
    return f"${value:,.2f} MXN"


def _fmt_num(value: float | int | None, decimals: int = 2) -> str:
    if value is None:
        return "—"
    return f"{value:,.{decimals}f}"


def _fmt_pct(value: float | None) -> str:
    if value is None:
        return "—"
    return f"{value:.1f}%"


def _set_col_width(cell, width_inches: float) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.get_or_add_tcW()
    tcW.set(qn("w:w"), str(int(width_inches * 1440)))
    tcW.set(qn("w:type"), "dxa")


def _set_cell_shading(cell, hex_color: str) -> None:
    """Aplica fondo de color a una celda usando lxml."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = etree.SubElement(tcPr, qn("w:shd"))
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#").upper())


def _set_cell_border(cell, color: str = "E5E7EB", size: int = 4) -> None:
    """Aplica bordes sutiles a una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = etree.SubElement(tcPr, qn("w:tcBorders"))
    for edge in ("top", "left", "bottom", "right"):
        edge_el = tcBorders.find(qn(f"w:{edge}"))
        if edge_el is None:
            edge_el = etree.SubElement(tcBorders, qn(f"w:{edge}"))
        edge_el.set(qn("w:val"), "single")
        edge_el.set(qn("w:sz"), str(size))
        edge_el.set(qn("w:space"), "0")
        edge_el.set(qn("w:color"), color)


def _add_styled_paragraph(
    doc: Document,
    text: str,
    *,
    bold: bool = False,
    size: int = 11,
    color=None,
    align=None,
    space_after=None,
) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color or _DARK_TEXT
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = _BRAND_BLUE
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(18)
    else:
        run.font.size = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(12)


def _style_header_cell(cell) -> None:
    """Pinta la celda de encabezado con fondo azul oscuro."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = etree.SubElement(tcPr, qn("w:shd"))
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "1E40AF")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p.runs:
        p.runs[0].font.color.rgb = _WHITE
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(10)
    _set_cell_border(cell, "1E40AF", 6)


def _style_data_cell(
    cell, align_right: bool = False, hex_bg: str | None = None, font_color=None
) -> None:
    p = cell.paragraphs[0]
    if align_right:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if p.runs:
        p.runs[0].font.size = Pt(10)
        if font_color:
            p.runs[0].font.color.rgb = font_color
    if hex_bg:
        _set_cell_shading(cell, hex_bg)
    _set_cell_border(cell, "E5E7EB", 4)


def _add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    *,
    zebra: bool = True,
    row_styles: list[dict | None] | None = None,
) -> None:
    """
    Crea una tabla con estilo profesional.
    row_styles: lista opcional de dicts con keys 'bg', 'color', 'align_right' por fila.
    """
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"
    table.autofit = False
    table.allow_autofit = False

    # Encabezados
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        _style_header_cell(cell)

    # Filas de datos
    for r_idx, row in enumerate(rows):
        is_even = r_idx % 2 == 0
        style_override = (row_styles or [None] * len(rows))[r_idx]
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx + 1, c_idx)
            cell.text = val

            bg = None
            color = None
            align_right = c_idx > 0

            if style_override:
                bg = style_override.get("bg")
                color = style_override.get("color")
                align_right = style_override.get("align_right", align_right)
            elif zebra and not is_even:
                bg = "F3F4F6"

            _style_data_cell(cell, align_right=align_right, hex_bg=bg, font_color=color)

    doc.add_paragraph()


class ReportService:
    def __init__(self, db: AsyncSession) -> None:
        self._db = db

    async def generate_ventas_docx(
        self,
        start_date: date | None,
        end_date: date | None,
        sections: list[str],
    ) -> bytes:
        svc = VentasService(self._db)
        selected = {s.lower().strip() for s in sections} & _SECTION_KEYS

        summary = (
            await svc.sales_summary(start_date=start_date, end_date=end_date)
            if "kpis" in selected
            else None
        )
        top_customers = (
            await svc.top_customers_by_sales(
                start_date=start_date, end_date=end_date, limit=15
            )
            if "clientes" in selected
            else []
        )
        gross_margin = (
            await svc.gross_margin_by_product(
                start_date=start_date, end_date=end_date, limit=15
            )
            if "margen" in selected
            else []
        )
        product_dist = (
            await svc.sales_distribution_by_product(
                start_date=start_date, end_date=end_date, limit=15
            )
            if "productos" in selected
            else []
        )
        pending_payments = (
            await svc.pending_payment_customers() if "pagos" in selected else []
        )
        at_risk = await svc.at_risk_customers() if "riesgo" in selected else []

        doc = Document()

        _set_page_margins(doc)
        _add_header_footer(doc, report_title="Reporte de Ventas")

        _build_cover(doc, start_date, end_date, report_title="Reporte de Ventas")

        if summary and "kpis" in selected:
            _build_kpis_section(doc, summary)

        if top_customers and "clientes" in selected:
            _build_top_customers_section(doc, top_customers)

        if product_dist and "productos" in selected:
            _build_product_distribution_section(doc, product_dist)

        if gross_margin and "margen" in selected:
            _build_gross_margin_section(doc, gross_margin)

        if pending_payments and "pagos" in selected:
            _build_pending_payments_section(doc, pending_payments)

        if at_risk and "riesgo" in selected:
            _build_at_risk_section(doc, at_risk)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    async def generate_ventas_csv(
        self,
        start_date: date | None,
        end_date: date | None,
        sections: list[str],
    ) -> bytes:
        """Genera un CSV multi-sección con los mismos datos que el DOCX."""
        svc = VentasService(self._db)
        selected = {s.lower().strip() for s in sections} & _SECTION_KEYS

        lines: list[str] = []

        if "kpis" in selected:
            summary = await svc.sales_summary(start_date=start_date, end_date=end_date)
            lines += [
                "# RESUMEN KPIs",
                "indicador,valor",
                f"Ventas Totales,{summary.total_sales}",
                f"Cotizaciones Aprobadas,{summary.approved_quotes}",
                f"Cotizaciones Pendientes,{summary.pending_quotes}",
                f"Cotizaciones Canceladas,{summary.cancelled_quotes}",
                f"Tasa de Conversión %,{summary.conversion_rate}",
                f"Margen Promedio %,{summary.average_margin_percent}",
                "",
            ]

        if "clientes" in selected:
            customers = await svc.top_customers_by_sales(
                start_date=start_date, end_date=end_date, limit=50
            )
            lines += [
                "# TOP CLIENTES POR VENTAS",
                "cliente,categoria,num_ventas,total_mxn,ticket_promedio_mxn",
            ]
            for c in customers:
                lines.append(
                    f"{_csv_esc(c.customer)},{_csv_esc(c.category or '')},"
                    f"{c.sale_count},{c.total_revenue},{c.average_ticket}"
                )
            lines.append("")

        if "productos" in selected:
            products = await svc.sales_distribution_by_product(
                start_date=start_date, end_date=end_date, limit=50
            )
            lines += [
                "# DISTRIBUCIÓN POR PRODUCTO",
                "producto,sku,unidades,revenue_mxn,porcentaje",
            ]
            for p in products:
                lines.append(
                    f"{_csv_esc(p.product)},{_csv_esc(p.sku or '')},"
                    f"{p.qty},{p.revenue},{p.percentage}"
                )
            lines.append("")

        if "margen" in selected:
            margins = await svc.gross_margin_by_product(
                start_date=start_date, end_date=end_date, limit=50
            )
            lines += [
                "# MARGEN BRUTO POR PRODUCTO",
                "producto,sku,ingresos_mxn,costo_mxn,margen_mxn,margen_pct",
            ]
            for m in margins:
                lines.append(
                    f"{_csv_esc(m.product)},{_csv_esc(m.sku or '')},"
                    f"{m.revenue},{m.cost},{m.gross_margin},{m.margin_percent or 0}"
                )
            lines.append("")

        if "pagos" in selected:
            payments = await svc.pending_payment_customers()
            lines += [
                "# PAGOS PENDIENTES",
                "cliente,tipo_cliente,pedidos,total_adeudado_mxn,dias_sin_pagar",
            ]
            for p in payments:
                lines.append(
                    f"{_csv_esc(p.customer_name)},{_csv_esc(p.tipo_cliente or '')},"
                    f"{p.num_pedidos},{p.total_adeudado},{p.dias_sin_pagar or ''}"
                )
            lines.append("")

        if "riesgo" in selected:
            at_risk = await svc.at_risk_customers()
            lines += [
                "# CLIENTES EN RIESGO DE ABANDONO",
                "cliente,riesgo,compras_ult_90_mxn,compras_90_previos_mxn,ultima_compra",
            ]
            for c in at_risk:
                ultima = c.ultima_compra.isoformat() if c.ultima_compra else ""
                lines.append(
                    f"{_csv_esc(c.customer_name)},{_csv_esc(c.riesgo_abandono)},"
                    f"{c.compras_ult_90},{c.compras_90_previos},{ultima}"
                )
            lines.append("")

        return "\n".join(lines).encode("utf-8-sig")

    async def generate_almacen_docx(
        self,
        start_date: date | None,
        end_date: date | None,
        sections: list[str],
    ) -> bytes:
        selected = {s.lower().strip() for s in sections} & _ALMACEN_SECTION_KEYS

        doc = Document()
        _set_page_margins(doc)
        _add_header_footer(doc, report_title="Reporte de Almacén")
        _build_cover(doc, start_date, end_date, report_title="Reporte de Almacén")

        kpis = await self._almacen_kpis() if "kpis" in selected else None
        top_value = (
            await self._almacen_top_value(limit=20) if "valor" in selected else []
        )
        low_stock = (
            await self._almacen_low_stock(limit=30) if "alertas" in selected else []
        )
        dormant = (
            await self._almacen_dormant(limit=30) if "dormidos" in selected else []
        )

        if kpis and "kpis" in selected:
            _build_almacen_kpis_section(doc, kpis)

        if top_value and "valor" in selected:
            _build_almacen_valor_section(doc, top_value)

        if low_stock and "alertas" in selected:
            _build_almacen_alertas_section(doc, low_stock)

        if dormant and "dormidos" in selected:
            _build_almacen_dormidos_section(doc, dormant)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    async def generate_almacen_csv(
        self,
        start_date: date | None,
        end_date: date | None,
        sections: list[str],
    ) -> bytes:
        selected = {s.lower().strip() for s in sections} & _ALMACEN_SECTION_KEYS

        lines: list[str] = []

        if "kpis" in selected:
            k = await self._almacen_kpis()
            lines += [
                "# RESUMEN KPIs (ALMACÉN)",
                "indicador,valor",
                f"SKUs totales,{k['total_skus']}",
                f"Valor total inventario MXN,{k['total_value_mxn']}",
                f"SKUs bajo mínimo,{k['low_stock_skus']}",
                f"SKUs sin stock,{k['no_stock_skus']}",
                f"SKUs sin movimiento >= 180 días,{k['dormant_skus']}",
                "",
            ]

        if "valor" in selected:
            items = await self._almacen_top_value(limit=200)
            lines += [
                "# TOP SKUs POR VALOR DE INVENTARIO",
                "producto,sku,qty_real,costo_unitario_mxn,valor_stock_mxn,dias_sin_mov,alerta",
            ]
            for it in items:
                lines.append(
                    f"{_csv_esc(it['product_name'])},{_csv_esc(it['sku'])},{it['real_qty']},"
                    f"{it['unit_cost_mxn']},{it['stock_total_cost_mxn']},{it['days_without_movement']},"
                    f"{_csv_esc(it['stock_alert'])}"
                )
            lines.append("")

        if "alertas" in selected:
            items = await self._almacen_low_stock(limit=200)
            lines += [
                "# ALERTAS DE STOCK (BAJO MÍNIMO / SIN STOCK)",
                "producto,sku,qty_real,costo_unitario_mxn,valor_stock_mxn,dias_sin_mov,alerta",
            ]
            for it in items:
                lines.append(
                    f"{_csv_esc(it['product_name'])},{_csv_esc(it['sku'])},{it['real_qty']},"
                    f"{it['unit_cost_mxn']},{it['stock_total_cost_mxn']},{it['days_without_movement']},"
                    f"{_csv_esc(it['stock_alert'])}"
                )
            lines.append("")

        if "dormidos" in selected:
            items = await self._almacen_dormant(limit=200)
            lines += [
                "# SKUs DORMIDOS (SIN MOVIMIENTO)",
                "producto,sku,dias_sin_mov,qty_real,valor_stock_mxn,alerta",
            ]
            for it in items:
                lines.append(
                    f"{_csv_esc(it['product_name'])},{_csv_esc(it['sku'])},{it['days_without_movement']},"
                    f"{it['real_qty']},{it['stock_total_cost_mxn']},{_csv_esc(it['stock_alert'])}"
                )
            lines.append("")

        header = [
            f"# Reporte de Almacén — {_period_label(start_date, end_date)}",
            f"# Generado: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}",
            "",
        ]
        return "\n".join(header + lines).encode("utf-8-sig")

    async def _almacen_kpis(self) -> dict[str, float | int]:
        total_skus = await self._db.scalar(select(func.count(InventoryItem.id)))
        total_value = await self._db.scalar(
            select(func.coalesce(func.sum(InventoryItem.stock_total_cost), 0))
        )

        low_stock = await self._db.scalar(
            select(func.count(InventoryItem.id)).where(
                InventoryItem.stock_alert.is_not(None),
                InventoryItem.stock_alert.ilike("%bajo%"),
            )
        )
        no_stock = await self._db.scalar(
            select(func.count(InventoryItem.id)).where(
                func.coalesce(InventoryItem.real_qty, 0) <= 0
            )
        )
        dormant = await self._db.scalar(
            select(func.count(InventoryItem.id)).where(
                InventoryItem.days_without_movement.is_not(None),
                InventoryItem.days_without_movement >= 180,
            )
        )

        return {
            "total_skus": int(total_skus or 0),
            "total_value_mxn": float(total_value or 0),
            "low_stock_skus": int(low_stock or 0),
            "no_stock_skus": int(no_stock or 0),
            "dormant_skus": int(dormant or 0),
        }

    async def _almacen_top_value(self, *, limit: int) -> list[dict[str, object]]:
        stmt = (
            select(
                Product.name,
                InventoryItem.internal_code,
                InventoryItem.external_product_id,
                InventoryItem.real_qty,
                InventoryItem.unit_cost,
                InventoryItem.stock_total_cost,
                InventoryItem.days_without_movement,
                InventoryItem.stock_alert,
            )
            .select_from(InventoryItem)
            .outerjoin(Product, InventoryItem.product_id == Product.id)
            .order_by(func.coalesce(InventoryItem.stock_total_cost, 0).desc())
            .limit(limit)
        )
        rows = (await self._db.execute(stmt)).all()
        return [_normalize_inventory_row(r) for r in rows]

    async def _almacen_low_stock(self, *, limit: int) -> list[dict[str, object]]:
        stmt = (
            select(
                Product.name,
                InventoryItem.internal_code,
                InventoryItem.external_product_id,
                InventoryItem.real_qty,
                InventoryItem.unit_cost,
                InventoryItem.stock_total_cost,
                InventoryItem.days_without_movement,
                InventoryItem.stock_alert,
            )
            .select_from(InventoryItem)
            .outerjoin(Product, InventoryItem.product_id == Product.id)
            .where(
                or_(
                    InventoryItem.stock_alert.ilike("%bajo%"),
                    func.coalesce(InventoryItem.real_qty, 0) <= 0,
                )
            )
            .order_by(func.coalesce(InventoryItem.real_qty, 0).asc())
            .limit(limit)
        )
        rows = (await self._db.execute(stmt)).all()
        return [_normalize_inventory_row(r) for r in rows]

    async def _almacen_dormant(self, *, limit: int) -> list[dict[str, object]]:
        stmt = (
            select(
                Product.name,
                InventoryItem.internal_code,
                InventoryItem.external_product_id,
                InventoryItem.real_qty,
                InventoryItem.unit_cost,
                InventoryItem.stock_total_cost,
                InventoryItem.days_without_movement,
                InventoryItem.stock_alert,
            )
            .select_from(InventoryItem)
            .outerjoin(Product, InventoryItem.product_id == Product.id)
            .where(
                InventoryItem.days_without_movement.is_not(None),
                InventoryItem.days_without_movement >= 180,
            )
            .order_by(InventoryItem.days_without_movement.desc())
            .limit(limit)
        )
        rows = (await self._db.execute(stmt)).all()
        return [_normalize_inventory_row(r) for r in rows]


def _csv_esc(value: str) -> str:
    if "," in value or '"' in value or "\n" in value:
        return '"' + value.replace('"', '""') + '"'
    return value


def _set_page_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)


def _add_header_footer(doc: Document, *, report_title: str) -> None:
    """Agrega encabezado y pie de página a todas las secciones del documento."""
    for section in doc.sections:
        # Header
        header = section.header
        header_para = (
            header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        )
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header_para.add_run(f"Nexus Ops RTB  ·  {report_title}")
        run.font.size = Pt(8)
        run.font.color.rgb = _GRAY
        run.font.italic = True
        header_para.paragraph_format.space_after = Pt(2)

        # Línea separadora en header
        p = header.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("─" * 60)
        run.font.size = Pt(6)
        run.font.color.rgb = RGBColor(0xE5, 0xE7, 0xEB)

        # Footer
        footer = section.footer
        footer_para = (
            footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        )
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.add_run("Documento confidencial  ·  Nexus Ops RTB")
        run.font.size = Pt(7)
        run.font.color.rgb = _GRAY


def _build_cover(
    doc: Document, start_date: date | None, end_date: date | None, *, report_title: str
) -> None:
    # Espaciado superior
    for _ in range(4):
        doc.add_paragraph()

    # Título principal
    title = doc.add_heading(report_title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title.runs:
        title.runs[0].font.color.rgb = _BRAND_BLUE
        title.runs[0].font.size = Pt(32)
        title.runs[0].font.bold = True
    title.paragraph_format.space_after = Pt(6)

    # Línea decorativa
    line_para = doc.add_paragraph()
    line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_para.paragraph_format.space_before = Pt(4)
    line_para.paragraph_format.space_after = Pt(12)
    line_run = line_para.add_run("━" * 20)
    line_run.font.color.rgb = _BRAND_BLUE
    line_run.font.size = Pt(10)

    # Subtítulo
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Nexus Ops RTB — Dashboard Operativo")
    sub_run.font.color.rgb = _GRAY
    sub_run.font.size = Pt(13)
    subtitle.paragraph_format.space_after = Pt(24)

    # Periodo en caja visual
    period_label = _period_label(start_date, end_date)
    period_para = doc.add_paragraph()
    period_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    period_para.paragraph_format.space_after = Pt(8)
    period_run = period_para.add_run(period_label)
    period_run.font.size = Pt(11)
    period_run.font.color.rgb = _DARK_TEXT
    period_run.bold = True

    # Fecha de generación
    gen_time = datetime.now(timezone.utc).strftime("%d/%m/%Y %H:%M UTC")
    gen = doc.add_paragraph()
    gen.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gen_run = gen.add_run(f"Generado: {gen_time}")
    gen_run.font.size = Pt(9)
    gen_run.font.color.rgb = _GRAY
    gen.paragraph_format.space_after = Pt(32)

    # Secciones incluidas placeholder (se puede expandir)
    doc.add_paragraph()

    doc.add_page_break()


def _period_label(start_date: date | None, end_date: date | None) -> str:
    if start_date and end_date:
        return f"Periodo: {start_date.strftime('%d/%m/%Y')} — {end_date.strftime('%d/%m/%Y')}"
    if start_date:
        return f"Desde: {start_date.strftime('%d/%m/%Y')}"
    if end_date:
        return f"Hasta: {end_date.strftime('%d/%m/%Y')}"
    return "Periodo: Todos los datos históricos"


def _build_kpis_section(doc: Document, summary) -> None:
    _add_heading(doc, "1. Resumen de KPIs", level=1)
    _add_styled_paragraph(
        doc,
        "Métricas clave del desempeño comercial en el periodo seleccionado.",
        color=_GRAY,
        size=10,
        space_after=12,
    )

    # Tabla de KPIs con fondo azul claro para los valores
    headers = ["Indicador", "Valor"]
    rows = [
        ["Ventas Totales", _fmt_mxn(summary.total_sales)],
        ["Cotizaciones Aprobadas", _fmt_num(summary.approved_quotes, 0)],
        ["Cotizaciones Pendientes", _fmt_num(summary.pending_quotes, 0)],
        ["Cotizaciones Canceladas", _fmt_num(summary.cancelled_quotes, 0)],
        ["Tasa de Conversión", _fmt_pct(summary.conversion_rate)],
        ["Margen Promedio", _fmt_pct(summary.average_margin_percent)],
    ]

    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"
    table.autofit = False
    table.allow_autofit = False

    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        _style_header_cell(cell)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx + 1, c_idx)
            cell.text = val
            if c_idx == 1:
                # Valor: fondo azul muy claro + negrita
                _set_cell_shading(cell, "DBEAFE")
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if p.runs:
                    p.runs[0].font.bold = True
                    p.runs[0].font.size = Pt(11)
                    p.runs[0].font.color.rgb = _BRAND_BLUE
            else:
                p = cell.paragraphs[0]
                if p.runs:
                    p.runs[0].font.size = Pt(10)
                    p.runs[0].bold = True
            _set_cell_border(cell, "E5E7EB", 4)

    doc.add_paragraph()


def _build_top_customers_section(doc: Document, customers: list) -> None:
    _add_heading(doc, "2. Top Clientes por Ventas", level=1)
    _add_styled_paragraph(
        doc,
        "Los siguientes clientes representan el mayor volumen de ventas en el periodo seleccionado.",
        color=_GRAY,
        size=10,
        space_after=12,
    )

    rows = [
        [
            str(i + 1),
            c.customer,
            c.category or "—",
            _fmt_num(c.sale_count, 0),
            _fmt_mxn(c.total_revenue),
            _fmt_mxn(c.average_ticket),
        ]
        for i, c in enumerate(customers)
    ]
    _add_table(
        doc,
        ["#", "Cliente", "Categoría", "Ventas", "Total MXN", "Ticket Prom."],
        rows,
        zebra=True,
    )


def _build_product_distribution_section(doc: Document, products: list) -> None:
    _add_heading(doc, "3. Distribución de Ventas por Producto", level=1)

    rows = [
        [
            p.product,
            p.sku or "—",
            _fmt_num(p.qty, 0),
            _fmt_mxn(p.revenue),
            _fmt_pct(p.percentage),
        ]
        for p in products
    ]
    _add_table(
        doc,
        ["Producto", "SKU", "Unidades", "Revenue MXN", "% del Total"],
        rows,
        zebra=True,
    )


def _build_gross_margin_section(doc: Document, products: list) -> None:
    _add_heading(doc, "4. Margen Bruto por Producto", level=1)
    _add_styled_paragraph(
        doc,
        "Comparación entre ingresos y costo de compra para los productos más relevantes.",
        color=_GRAY,
        size=10,
        space_after=12,
    )

    row_styles: list[dict | None] = []
    for p in products:
        margin_pct = p.margin_percent
        style: dict | None = None
        if margin_pct is not None:
            if margin_pct < 0:
                style = {"bg": _RED_BG, "color": _RED}
            elif margin_pct >= 30:
                style = {"bg": _GREEN_BG, "color": _GREEN}
        row_styles.append(style)

    rows = [
        [
            p.product,
            p.sku or "—",
            _fmt_mxn(p.revenue),
            _fmt_mxn(p.cost),
            _fmt_mxn(p.gross_margin),
            _fmt_pct(p.margin_percent),
        ]
        for p in products
    ]
    _add_table(
        doc,
        ["Producto", "SKU", "Ingresos", "Costo", "Margen MXN", "% Margen"],
        rows,
        zebra=True,
        row_styles=row_styles,
    )


def _build_pending_payments_section(doc: Document, payments: list) -> None:
    _add_heading(doc, "5. Pagos Pendientes por Cliente", level=1)
    _add_styled_paragraph(
        doc,
        "Clientes con cotizaciones aprobadas aún sin registrar pago completo.",
        color=_GRAY,
        size=10,
        space_after=12,
    )

    row_styles: list[dict | None] = []
    for p in payments:
        dias = p.dias_sin_pagar
        style: dict | None = None
        if dias is not None:
            if dias >= 60:
                style = {"bg": _RED_BG, "color": _RED}
            elif dias >= 30:
                style = {"bg": _ORANGE_BG, "color": _ORANGE}
        row_styles.append(style)

    rows = [
        [
            p.customer_name,
            p.tipo_cliente or "—",
            _fmt_num(p.num_pedidos, 0),
            _fmt_mxn(p.total_adeudado),
            str(p.dias_sin_pagar) + " días" if p.dias_sin_pagar is not None else "—",
        ]
        for p in payments
    ]
    _add_table(
        doc,
        ["Cliente", "Tipo", "Pedidos", "Total Adeudado MXN", "Sin pagar"],
        rows,
        zebra=True,
        row_styles=row_styles,
    )


def _build_at_risk_section(doc: Document, customers: list) -> None:
    _add_heading(doc, "6. Clientes en Riesgo de Abandono", level=1)
    _add_styled_paragraph(
        doc,
        "Clientes con caída significativa en compras comparando los últimos 90 días contra los 90 días previos.",
        color=_GRAY,
        size=10,
        space_after=12,
    )

    row_styles: list[dict | None] = []
    for c in customers:
        riesgo = (c.riesgo_abandono or "").lower().strip()
        style: dict | None = None
        if riesgo == "crítico":
            style = {"bg": _RED_BG, "color": _RED}
        elif riesgo == "alto":
            style = {"bg": _ORANGE_BG, "color": _ORANGE}
        elif riesgo == "medio":
            style = {"bg": _YELLOW_BG, "color": _YELLOW}
        row_styles.append(style)

    rows = [
        [
            c.customer_name,
            c.riesgo_abandono or "—",
            _fmt_mxn(c.compras_ult_90),
            _fmt_mxn(c.compras_90_previos),
            c.ultima_compra.strftime("%d/%m/%Y") if c.ultima_compra else "—",
        ]
        for c in customers
    ]
    _add_table(
        doc,
        ["Cliente", "Riesgo", "Últimos 90 días", "90 días previos", "Última compra"],
        rows,
        zebra=True,
        row_styles=row_styles,
    )


def _normalize_inventory_row(row: tuple) -> dict[str, object]:
    (
        product_name,
        internal_code,
        external_id,
        real_qty,
        unit_cost,
        total_cost,
        days_wo,
        stock_alert,
    ) = row
    sku = internal_code or external_id or "—"
    return {
        "product_name": product_name or "—",
        "sku": sku,
        "real_qty": float(real_qty or 0),
        "unit_cost_mxn": float(unit_cost or 0),
        "stock_total_cost_mxn": float(total_cost or 0),
        "days_without_movement": int(days_wo or 0),
        "stock_alert": str(stock_alert or ""),
    }


def _build_almacen_kpis_section(doc: Document, kpis: dict[str, float | int]) -> None:
    _add_heading(doc, "1. Resumen de KPIs (Almacén)", level=1)
    _add_styled_paragraph(
        doc,
        "Indicadores operativos del inventario actual.",
        color=_GRAY,
        size=10,
        space_after=12,
    )

    rows = [
        ["SKUs totales", _fmt_num(kpis["total_skus"], 0)],
        ["Valor total inventario", _fmt_mxn(float(kpis["total_value_mxn"]))],
        ["SKUs bajo mínimo", _fmt_num(kpis["low_stock_skus"], 0)],
        ["SKUs sin stock", _fmt_num(kpis["no_stock_skus"], 0)],
        ["SKUs sin movimiento ≥ 180 días", _fmt_num(kpis["dormant_skus"], 0)],
    ]
    _add_table(doc, ["Indicador", "Valor"], rows, zebra=True)


def _build_almacen_valor_section(doc: Document, items: list[dict[str, object]]) -> None:
    _add_heading(doc, "2. Top SKUs por Valor de Inventario", level=1)
    _add_styled_paragraph(
        doc,
        "Productos con mayor valor de inventario (costo total en stock).",
        color=_GRAY,
        size=10,
        space_after=12,
    )

    rows = [
        [
            str(it["product_name"]),
            str(it["sku"]),
            _fmt_num(float(it["real_qty"]), 0),
            _fmt_mxn(float(it["unit_cost_mxn"])),
            _fmt_mxn(float(it["stock_total_cost_mxn"])),
            _fmt_num(int(it["days_without_movement"]), 0),
        ]
        for it in items
    ]
    _add_table(
        doc,
        ["Producto", "SKU", "Qty real", "Costo unit.", "Valor stock", "Días sin mov"],
        rows,
        zebra=True,
    )


def _build_almacen_alertas_section(
    doc: Document, items: list[dict[str, object]]
) -> None:
    _add_heading(doc, "3. Alertas de Stock", level=1)
    _add_styled_paragraph(
        doc,
        "SKUs con bajo mínimo o sin stock para priorizar reposición.",
        color=_GRAY,
        size=10,
        space_after=12,
    )

    row_styles: list[dict | None] = []
    for it in items:
        style: dict | None = None
        alert = str(it["stock_alert"]).lower()
        if "bajo" in alert or float(it["real_qty"]) <= 0:
            style = {"bg": _RED_BG, "color": _RED}
        row_styles.append(style)

    rows = [
        [
            str(it["product_name"]),
            str(it["sku"]),
            _fmt_num(float(it["real_qty"]), 0),
            _fmt_mxn(float(it["stock_total_cost_mxn"])),
            str(it["stock_alert"]) or "—",
        ]
        for it in items
    ]
    _add_table(
        doc,
        ["Producto", "SKU", "Qty real", "Valor stock", "Alerta"],
        rows,
        zebra=True,
        row_styles=row_styles,
    )


def _build_almacen_dormidos_section(
    doc: Document, items: list[dict[str, object]]
) -> None:
    _add_heading(doc, "4. SKUs Dormidos (Sin Movimiento)", level=1)
    _add_styled_paragraph(
        doc,
        "Productos con más días sin movimiento para análisis de obsolescencia/rotación.",
        color=_GRAY,
        size=10,
        space_after=12,
    )

    rows = [
        [
            str(it["product_name"]),
            str(it["sku"]),
            _fmt_num(int(it["days_without_movement"]), 0),
            _fmt_num(float(it["real_qty"]), 0),
            _fmt_mxn(float(it["stock_total_cost_mxn"])),
        ]
        for it in items
    ]
    _add_table(
        doc,
        ["Producto", "SKU", "Días sin mov", "Qty real", "Valor stock"],
        rows,
        zebra=True,
    )
